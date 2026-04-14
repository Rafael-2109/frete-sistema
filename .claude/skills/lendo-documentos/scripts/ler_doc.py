#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Leitura de Documentos — Word, CNAB (.ret/.rem/.cnab), OFX

Skill: lendo-documentos (complementa lendo-arquivos para Excel/CSV)

Formatos:
- .docx: paragraphs + tables via python-docx
- .ret: CNAB400 retorno via Cnab400ParserService (layout BMP 274)
- .rem: estrutural basico (header/detalhe/trailer sem extracao de campos)
- .cnab: auto-deteccao (assume formato retorno se tipo 0/1/9 presentes)
- .ofx: Open Financial Exchange via parsear_ofx (latin-1)

Reusa parsers validados em producao de app/financeiro/services (standalone,
sem dependencia de Flask app context).

Dependencias:
- python-docx (pip install python-docx==1.1.2)  — .docx
- stdlib apenas                                    — .ret/.rem/.cnab/.ofx
"""
import sys
import os
import json
import argparse
import tempfile
import traceback
from datetime import datetime, date
from decimal import Decimal

# Adicionar path do projeto para imports de app.financeiro.services
sys.path.insert(
    0,
    os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../..')),
)


def decimal_default(obj):
    """Serializa tipos especiais (Decimal, date, datetime, bytes) para JSON."""
    if isinstance(obj, Decimal):
        return float(obj)
    if isinstance(obj, (datetime, date)):
        return obj.isoformat()
    if isinstance(obj, bytes):
        # Parsers podem retornar bytes nao-decodificados em campos raros
        try:
            return obj.decode('latin-1')
        except Exception:
            return obj.hex()
    if hasattr(obj, 'isoformat'):
        return obj.isoformat()
    if obj != obj:  # NaN check
        return None
    raise TypeError(f"Object of type {type(obj)} is not JSON serializable")


def url_para_caminho(url, user_id=None):
    """
    Converte URL do agente para caminho local do arquivo.

    Formatos aceitos:
      - /agente/api/files/{session_id}/{filename}
      - /tmp/agente_files/{user_id}/{session_id}/{filename}
      - Caminho absoluto direto

    Args:
        url: URL ou caminho absoluto
        user_id: se fornecido, restringe o scan a ESTE usuario apenas
            (previne cross-user leak). Quando None, faz scan em todos
            os subdiretorios — modo legacy para CLI direto.

    Copia do padrao de lendo-arquivos/scripts/ler.py para consistencia.
    """
    # Basic sanitization: rejeita tentativas obvias de path traversal
    if '..' in url or '\x00' in url:
        return url  # deixa existir check falhar depois

    if (
        url.startswith('/')
        and not url.startswith('/agente')
        and os.path.exists(url)
    ):
        # Caminho absoluto real — ja foi resolvido pelo chamador
        return url

    parts = url.strip('/').split('/')
    if (
        len(parts) >= 4
        and parts[0] == 'agente'
        and parts[1] == 'api'
        and parts[2] == 'files'
    ):
        session_id = parts[3]
        filename = parts[4] if len(parts) > 4 else parts[3]

        # Reject session_id/filename com path traversal
        if '..' in session_id or '/' in session_id or '..' in filename:
            return url

        base_folder = os.path.join(tempfile.gettempdir(), 'agente_files')

        # Modo restrito: so tenta caminho do user fornecido
        if user_id is not None:
            filepath_user = os.path.join(
                base_folder, str(user_id), session_id, filename
            )
            if os.path.exists(filepath_user):
                return filepath_user
            # Fallback: caminho direto (skills-generated files sem user)
            filepath_direct = os.path.join(base_folder, session_id, filename)
            return filepath_direct

        # Modo legacy (user_id nao fornecido): tenta caminho direto primeiro
        filepath_direct = os.path.join(base_folder, session_id, filename)
        if os.path.exists(filepath_direct):
            return filepath_direct

        # Legacy scan em subdirs user — so quando CLI explicitamente sem user_id
        if os.path.exists(base_folder):
            for user_dir in os.listdir(base_folder):
                user_path = os.path.join(base_folder, user_dir)
                if os.path.isdir(user_path):
                    filepath_user = os.path.join(
                        user_path, session_id, filename
                    )
                    if os.path.exists(filepath_user):
                        return filepath_user

        return filepath_direct

    return url


# =============================================================================
# Parsers por tipo
# =============================================================================

def ler_docx(filepath, limite=1000):
    """
    Le arquivo .docx via python-docx.

    Retorna: paragraphs (strip vazios), tables (lista de listas de strings),
    metadata (titulo, autor, created, modified, contagens).
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx nao instalado. Execute: pip install python-docx==1.1.2"
        )

    doc = Document(filepath)

    # Metadata de core_properties
    core = doc.core_properties
    metadata = {
        'titulo': core.title or None,
        'autor': core.author or None,
        'created': core.created.isoformat() if core.created else None,
        'modified': core.modified.isoformat() if core.modified else None,
        'paragraphs_count': len(doc.paragraphs),
        'tables_count': len(doc.tables),
    }

    # Paragraphs (strip vazios, aplicar limite)
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)
        if limite and len(paragraphs) >= limite:
            break

    # Tables (todas, sem limite — sao geralmente poucas)
    tables = []
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)

    return {
        'metadata': metadata,
        'paragraphs': paragraphs,
        'tables': tables,
        'total_paragraphs': len(doc.paragraphs),
        'paragraphs_retornados': len(paragraphs),
    }


def ler_cnab_retorno(filepath, limite=1000, offset=0):
    """
    Le arquivo CNAB400 retorno (.ret/.cnab) via Cnab400ParserService.

    Layout: BMP 274 (Banco Money Plus). Outros bancos podem divergir.
    Encoding: latin-1 (ISO-8859-1).
    """
    try:
        from app.financeiro.services.cnab400_parser_service import (
            Cnab400ParserService,
        )
    except ImportError as e:
        raise ImportError(f"Cnab400ParserService nao disponivel: {e}")

    with open(filepath, 'rb') as f:
        raw = f.read()
    try:
        conteudo = raw.decode('latin-1')
    except UnicodeDecodeError:
        conteudo = raw.decode('utf-8', errors='replace')

    parser = Cnab400ParserService()
    resultado = parser.parse_arquivo(conteudo)

    header = resultado.get('header') or {}
    trailer = resultado.get('trailer') or {}
    detalhes_raw = resultado.get('detalhes') or []
    erros = resultado.get('erros') or []

    # Paginacao
    total_detalhes = len(detalhes_raw)
    if limite:
        detalhes_page = detalhes_raw[offset:offset + limite]
    else:
        detalhes_page = detalhes_raw[offset:]

    # Remover 'linha_original' (poluente no JSON, ja temos os campos parseados)
    detalhes_limpos = [
        {k: v for k, v in d.items() if k != 'linha_original'}
        for d in detalhes_page
    ]

    return {
        'banco': header.get('nome_banco'),
        'codigo_banco': header.get('codigo_banco'),
        'empresa': header.get('nome_empresa'),
        'cnpj_empresa': header.get('cnpj_empresa'),
        'data_arquivo': header.get('data_arquivo'),
        'total_detalhes': total_detalhes,
        'detalhes_retornados': len(detalhes_limpos),
        'qtd_titulos_trailer': trailer.get('qtd_titulos'),
        'valor_total_trailer': trailer.get('valor_total'),
        'header': header,
        'trailer': trailer,
        'detalhes': detalhes_limpos,
        'erros': erros,
    }


def ler_cnab_remessa(filepath, limite=1000, offset=0):
    """
    Le .rem (remessa) com parser estrutural basico.

    NAO extrai campos (layout varia muito por banco — Santander, Itau, BB,
    Caixa, Bradesco tem posicoes diferentes). Retorna linhas brutas agrupadas
    por tipo de registro (0=header, 1=detalhe, 9=trailer).

    Para extracao de campos, o agente deve fazer parse heuristico a partir do
    conteudo cru ou invocar um parser especifico do banco.
    """
    with open(filepath, 'rb') as f:
        raw = f.read()
    try:
        conteudo = raw.decode('latin-1')
    except UnicodeDecodeError:
        conteudo = raw.decode('utf-8', errors='replace')

    linhas = conteudo.strip().split('\n')
    header = None
    trailer = None
    detalhes_raw = []

    for i, linha in enumerate(linhas, 1):
        linha = linha.rstrip('\r\n')
        if not linha:
            continue
        tipo = linha[0] if linha else ''
        registro = {
            'numero_linha': i,
            'tipo': tipo,
            'tamanho': len(linha),
            'conteudo': linha,  # linha inteira (posicional, cru)
        }
        if tipo == '0':
            header = registro
        elif tipo == '9':
            trailer = registro
        else:
            detalhes_raw.append(registro)

    total = len(detalhes_raw)
    if limite:
        detalhes_page = detalhes_raw[offset:offset + limite]
    else:
        detalhes_page = detalhes_raw[offset:]

    return {
        'formato': 'CNAB400 remessa (estrutural, sem extracao de campos)',
        'total_detalhes': total,
        'detalhes_retornados': len(detalhes_page),
        'header': header,
        'trailer': trailer,
        'detalhes': detalhes_page,
    }


def ler_ofx(filepath, limite=1000, offset=0):
    """
    Le arquivo .ofx via parsear_ofx (standalone).

    Formato: SGML OFX (brasileiro). Encoding tipico: latin-1.
    Retorna: acctid, dtstart, dtend, transacoes (paginadas).
    """
    try:
        from app.financeiro.services.ofx_parser_service import parsear_ofx
    except ImportError as e:
        raise ImportError(f"ofx_parser_service nao disponivel: {e}")

    with open(filepath, 'rb') as f:
        raw = f.read()

    resultado = parsear_ofx(raw)
    transacoes_raw = resultado.get('transacoes', []) or []
    total = len(transacoes_raw)
    if limite:
        transacoes_page = transacoes_raw[offset:offset + limite]
    else:
        transacoes_page = transacoes_raw[offset:]

    return {
        'acctid': resultado.get('acctid'),
        'dtstart': resultado.get('dtstart'),
        'dtend': resultado.get('dtend'),
        'total_transacoes': total,
        'transacoes_retornadas': len(transacoes_page),
        'transacoes': transacoes_page,
    }


# =============================================================================
# Main
# =============================================================================

def _detectar_tipo(extensao: str, tipo_arg: str) -> str:
    """Detecta tipo real a partir da extensao ou forca via --tipo."""
    if tipo_arg != 'auto':
        return tipo_arg
    if extensao == 'docx':
        return 'docx'
    if extensao in ('ret', 'cnab'):
        return 'cnab'
    if extensao == 'rem':
        return 'rem'
    if extensao == 'ofx':
        return 'ofx'
    return ''


def main():
    parser = argparse.ArgumentParser(
        description='Le Word (.docx), CNAB (.ret/.rem/.cnab) e OFX; retorna JSON.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemplos:
  python ler_doc.py --url "/agente/api/files/default/abc_contrato.docx"
  python ler_doc.py --url "/agente/api/files/default/abc_retorno.ret"
  python ler_doc.py --url "/agente/api/files/default/abc_extrato.ofx"
  python ler_doc.py --url "..." --limite 100 --offset 200
        """,
    )
    parser.add_argument('--url', required=True,
                        help='URL do anexo ou caminho absoluto')
    parser.add_argument(
        '--tipo',
        choices=['auto', 'docx', 'cnab', 'ret', 'rem', 'ofx'],
        default='auto',
        help='Forca tipo (default auto por extensao)',
    )
    parser.add_argument('--limite', type=int, default=1000,
                        help='Max linhas/transacoes retornadas (default 1000)')
    parser.add_argument('--offset', type=int, default=0,
                        help='Offset para paginacao (default 0)')
    parser.add_argument(
        '--user-id',
        type=int,
        default=None,
        help=(
            'User ID para restringir scan cross-user '
            '(previne leak entre usuarios). Quando omitido, '
            'usa modo legacy de scan completo.'
        ),
    )

    args = parser.parse_args()

    resultado = {
        'sucesso': False,
        'arquivo': None,
        'dados': None,
        'resumo': '',
    }

    try:
        filepath = url_para_caminho(args.url, user_id=args.user_id)
        if not os.path.exists(filepath):
            resultado['erro'] = f'Arquivo nao encontrado: {filepath}'
            resultado['mensagem'] = 'Verifique a URL do anexo'
            resultado['url_recebida'] = args.url
            print(json.dumps(
                resultado, ensure_ascii=False, indent=2,
                default=decimal_default,
            ))
            return

        tamanho = os.path.getsize(filepath)
        filename = os.path.basename(filepath)
        extensao = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

        tipo = _detectar_tipo(extensao, args.tipo)
        if not tipo:
            resultado['erro'] = f'Extensao sem suporte: .{extensao}'
            resultado['mensagem'] = 'Suportados: .docx, .ret, .rem, .cnab, .ofx'
            print(json.dumps(
                resultado, ensure_ascii=False, indent=2,
                default=decimal_default,
            ))
            return

        # Normaliza 'ret' → 'cnab' (mesmo parser)
        if tipo == 'ret':
            tipo = 'cnab'

        # Despacho por tipo
        if tipo == 'docx':
            dados = ler_docx(filepath, args.limite)
            resultado['dados'] = dados
            total_p = dados.get('total_paragraphs', 0)
            autor = (dados.get('metadata') or {}).get('autor') or 'N/A'
            resultado['resumo'] = (
                f"Word com {total_p} paragraphs e "
                f"{len(dados.get('tables', []))} tabelas. Autor: {autor}"
            )

        elif tipo == 'cnab':
            dados = ler_cnab_retorno(filepath, args.limite, args.offset)
            resultado['dados'] = dados
            banco = dados.get('banco') or '?'
            total = dados.get('total_detalhes', 0)
            valor = dados.get('valor_total_trailer') or 0
            resultado['resumo'] = (
                f"CNAB {banco}: {total} titulos, "
                f"valor total trailer R$ {valor:.2f}"
            )

        elif tipo == 'rem':
            dados = ler_cnab_remessa(filepath, args.limite, args.offset)
            resultado['dados'] = dados
            total = dados.get('total_detalhes', 0)
            resultado['resumo'] = (
                f"CNAB remessa estrutural: {total} registros "
                f"(sem extracao de campos — layout varia por banco)"
            )

        elif tipo == 'ofx':
            dados = ler_ofx(filepath, args.limite, args.offset)
            resultado['dados'] = dados
            acctid = dados.get('acctid') or '?'
            total_t = dados.get('total_transacoes', 0)
            resultado['resumo'] = (
                f"OFX conta {acctid}: {total_t} transacoes "
                f"({dados.get('dtstart')} a {dados.get('dtend')})"
            )

        resultado['sucesso'] = True
        resultado['arquivo'] = {
            'nome': filename,
            'tipo': tipo,
            'tamanho': tamanho,
            'tamanho_formatado': (
                f"{tamanho / 1024:.1f} KB" if tamanho < 1024 * 1024
                else f"{tamanho / (1024 * 1024):.1f} MB"
            ),
        }

    except ImportError as e:
        resultado['erro'] = f'Dependencia nao instalada: {e}'
        resultado['mensagem'] = (
            'Para .docx: pip install python-docx==1.1.2. '
            'CNAB/OFX usam parsers do proprio projeto (sem pip).'
        )
    except Exception as e:
        resultado['erro'] = str(e)
        resultado['traceback'] = traceback.format_exc()
        resultado['mensagem'] = 'Erro ao processar arquivo'

    print(json.dumps(
        resultado, ensure_ascii=False, indent=2, default=decimal_default,
    ))


if __name__ == '__main__':
    main()
